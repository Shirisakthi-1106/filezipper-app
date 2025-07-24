import heapq
import pickle
import numpy as np
from collections import defaultdict
from PIL import Image
import docx
import os

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None  # For PDF reading (text-based)


class HuffmanCoding:
    class Node:
        def __init__(self, symbol, freq):
            self.symbol = symbol
            self.freq = freq
            self.left = None
            self.right = None

        def __lt__(self, other):
            return self.freq < other.freq

    def __init__(self, file_path):
        self.file_path = file_path
        self.file_type = os.path.splitext(file_path)[1].lower().replace('.', '')
        self.heap = []
        self.codes = {}
        self.reverse_mapping = {}

    # ---------- Common Huffman Logic ----------
    def _build_frequency_dict(self, data):
        freq = defaultdict(int)
        if self.file_type in ['jpg', 'jpeg', 'png','bmp']:
            for pixel in map(tuple, data):  # pixels
                freq[pixel] += 1
        else:
            for char in data:  # text
                freq[char] += 1
        return freq

    def _build_heap(self, freq_dict):
        for symbol, freq in freq_dict.items():
            heapq.heappush(self.heap, self.Node(symbol, freq))

    def _merge_nodes(self):
        while len(self.heap) > 1:
            node1 = heapq.heappop(self.heap)
            node2 = heapq.heappop(self.heap)
            merged = self.Node(None, node1.freq + node2.freq)
            merged.left = node1
            merged.right = node2
            heapq.heappush(self.heap, merged)

    def _build_codes_helper(self, node, code=""):
        if node:
            if node.symbol is not None:
                self.codes[node.symbol] = code
                self.reverse_mapping[code] = node.symbol
            self._build_codes_helper(node.left, code + "0")
            self._build_codes_helper(node.right, code + "1")

    def _build_codes(self):
        root = heapq.heappop(self.heap)
        self._build_codes_helper(root)

    def _get_encoded_data(self, data):
        if self.file_type in ['jpg', 'jpeg', 'png','bmp']:
            return "".join(self.codes[tuple(pixel)] for pixel in data)
        else:
            return "".join(self.codes[char] for char in data)

    def _pad_encoded_data(self, encoded_data):
        extra_padding = 8 - len(encoded_data) % 8
        padded_info = "{0:08b}".format(extra_padding)
        return padded_info + encoded_data + "0" * extra_padding, extra_padding

    def _get_byte_array(self, padded_data):
        return bytearray(int(padded_data[i:i + 8], 2) for i in range(0, len(padded_data), 8))

    def _remove_padding(self, padded_data):
        padding = int(padded_data[:8], 2)
        return padded_data[8:-padding]

    def _decode_data(self, bit_string):
        current_code = ""
        output = []
        for bit in bit_string:
            current_code += bit
            if current_code in self.reverse_mapping:
                output.append(self.reverse_mapping[current_code])
                current_code = ""
        return output

    # ---------- File Readers ----------
    def _read_text(self):
        with open(self.file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _read_docx(self):
        doc = docx.Document(self.file_path)
        return "\n".join(para.text for para in doc.paragraphs)

    def _read_pdf(self):
        if PyPDF2 is None:
            raise ImportError("Install PyPDF2 to read PDF files.")
        reader = PyPDF2.PdfReader(self.file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _read_image(self):
        img = Image.open(self.file_path).convert("RGB")
        array = np.array(img)
        return array.reshape(-1, 3), array.shape

    # ---------- File Writers ----------
    def _write_text(self, data, out_path):
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(data)

    def _write_docx(self, data, out_path):
        doc = docx.Document()
        for line in data.split('\n'):
            doc.add_paragraph(line)
        doc.save(out_path)

    def _write_image(self, pixels, shape, out_path):
        array = np.array(pixels, dtype=np.uint8).reshape(shape)
        img = Image.fromarray(array, 'RGB')
        img.save(out_path)

    # ---------- Public Compression ----------
    def compress(self):
        try:
            if self.file_type == 'txt':
                data = self._read_text()
            elif self.file_type == 'docx':
                data = self._read_docx()
            elif self.file_type == 'pdf':
                data = self._read_pdf()
            elif self.file_type in ['jpg', 'jpeg', 'png','bmp']:
                data, self.image_shape = self._read_image()
            else:
                raise ValueError(f"Unsupported file type: {self.file_type}")
        except Exception as e:
            print(f"Read error: {e}")
            return None

        freq_dict = self._build_frequency_dict(data)
        self._build_heap(freq_dict)
        self._merge_nodes()
        self._build_codes()

        encoded_data = self._get_encoded_data(data)
        padded_data, padding = self._pad_encoded_data(encoded_data)
        byte_array = self._get_byte_array(padded_data)

        output_path = self.file_path + ".bin"

        with open(output_path, 'wb') as f:
            pickle.dump({
                'bytes': byte_array,
                'reverse_mapping': self.reverse_mapping,
                'padding': padding,
                'type': self.file_type,
                'shape': getattr(self, 'image_shape', None)
            }, f)

        print(f"Compressed to: {output_path}")
        return output_path

    # ---------- Public Decompression ----------
    def decompress(self):
        try:
            with open(self.file_path, 'rb') as f:
                data = pickle.load(f)

            self.reverse_mapping = data['reverse_mapping']
            byte_array = data['bytes']
            padding = data['padding']
            file_type = data['type']
            shape = data.get('shape', None)

            bit_string = ''.join(f"{byte:08b}" for byte in byte_array)
            bit_string = self._remove_padding(bit_string)
            decoded = self._decode_data(bit_string)
        except Exception as e:
            print(f"Decompression error: {e}")
            return None

        output_path = self.file_path.replace(".bin", f"_decompressed.{file_type}")

        try:
            if file_type in ['txt', 'pdf']:
                self._write_text("".join(decoded), output_path)
            elif file_type == 'docx':
                self._write_docx("".join(decoded), output_path)
            elif file_type in ['jpg', 'jpeg', 'png','bmp']:
                self._write_image(decoded, shape, output_path)
            else:
                raise ValueError(f"Unsupported file type during decompression: {file_type}")
        except Exception as e:
            print(f"Write error: {e}")
            return None

        print(f"Decompressed to: {output_path}")
        return output_path
